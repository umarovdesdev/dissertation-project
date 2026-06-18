"""Combine the GOST front matter into ONE document (EN + KZ) as .docx (+ .pdf).

Order follows the real IITU sample dissertations:
    TITLE PAGE → CONTENTS → NORMATIVE REFERENCES →
    DESIGNATIONS AND ABBREVIATIONS → DEFINITIONS

The whole bundle is a single Word section so page numbering is continuous and
the title page (page 1) carries no number — exactly as in the council samples
(contents = p. 2, normative = p. 3, …). Each section starts on a new page.

The CONTENTS page numbers are read from the already-assembled manuscript
(defense/docs/DISSERTATION_{EN,KZ}_GOST_<date>.docx) via MS Word, identical to
the standalone table of contents; entries not yet in the manuscript get an
em-dash, so the bundle stays honest until final assembly.

Sources : thesis/output/{titlepage,contents,normative_references,abbreviations,
          definitions}_{en,kz}.md  (titlepage values live in build_title.FIELDS)
Output  : defense/docs/FRONT_MATTER_{EN,KZ}_GOST_<date>.docx (+ .pdf)

Usage:
    python build_frontmatter_bundle.py [--date YYYY-MM-DD] [--no-pdf]
"""
from __future__ import annotations

import argparse
import importlib.util
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

_HERE = Path(__file__).resolve().parent
ROOT = _HERE
while ROOT.parent != ROOT and not (ROOT / "defense").is_dir():
    ROOT = ROOT.parent


def _load(name: str):
    spec = importlib.util.spec_from_file_location(name, _HERE / f"{name}.py")
    mod = importlib.util.module_from_spec(spec)
    spec.loader.exec_module(mod)
    return mod


md2gost = _load("md2gost")
build_title = _load("build_title")
build_toc = _load("build_toc")
build_frontmatter = _load("build_frontmatter")


def render_simple_md(doc, text: str) -> None:
    """Render heading + paragraphs (normative references / definitions)."""
    text = md2gost.strip_version_markers(text)
    block: list[str] = []

    def flush():
        if not block:
            return
        joined = " ".join(block).strip()
        block.clear()
        if joined.startswith("# "):
            md2gost._heading(doc, joined[2:].strip(), 1)
        else:
            md2gost._body(doc, joined)

    for raw in text.splitlines():
        s = raw.strip()
        if not s:
            flush()
            continue
        if s.startswith("# "):  # heading is always its own block
            flush()
            block.append(s)
            flush()
            continue
        block.append(s)
    flush()


def render_contents(doc, md_path: Path, num, front) -> None:
    for level, kind, text in build_toc.parse_md(md_path):
        if kind == "title":
            h = doc.add_paragraph()
            h.alignment = WD_ALIGN_PARAGRAPH.CENTER
            h.paragraph_format.space_after = Pt(12)
            r = h.add_run(text)
            md2gost._set_cell_font(r, bold=True)
            r.font.size = Pt(16)
        else:
            build_toc.add_entry(
                doc, text, build_toc.resolve_page(text, num, front),
                bold=(kind == "main"), indent_mm=build_toc.INDENT_MM[level],
            )


def build(lang: str, num, front, out_docx: Path) -> None:
    src = ROOT / "thesis/output"
    suffix = lang  # 'en' / 'kz'
    doc = Document()
    md2gost._configure_styles(doc)
    md2gost._configure_page(doc)

    # 1. Title page (page 1, unnumbered)
    build_title.populate(doc, lang)

    # 2. Contents
    doc.add_page_break()
    render_contents(doc, src / f"contents_{suffix}.md", num, front)

    # 3. Normative references
    doc.add_page_break()
    render_simple_md(doc, (src / f"normative_references_{suffix}.md").read_text(encoding="utf-8"))

    # 4. Designations and abbreviations
    doc.add_page_break()
    build_frontmatter.render_abbrev_into(
        doc, md2gost.strip_version_markers((src / f"abbreviations_{suffix}.md").read_text(encoding="utf-8"))
    )

    # 5. Definitions
    doc.add_page_break()
    render_simple_md(doc, (src / f"definitions_{suffix}.md").read_text(encoding="utf-8"))

    md2gost._add_page_numbers(doc)
    out_docx.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_docx))
    print("[docx]", out_docx.name)


def main() -> None:
    ap = argparse.ArgumentParser(description="Combine GOST front matter (EN+KZ)")
    ap.add_argument("--date", default="2026-06-17")
    ap.add_argument("--no-pdf", action="store_true")
    args = ap.parse_args()

    docs = ROOT / "defense/docs"
    jobs = [
        ("en", docs / f"DISSERTATION_EN_GOST_{args.date}.docx",
         docs / f"FRONT_MATTER_EN_GOST_{args.date}.docx"),
        ("kz", docs / f"DISSERTATION_KZ_GOST_{args.date}.docx",
         docs / f"FRONT_MATTER_KZ_GOST_{args.date}.docx"),
    ]

    import win32com.client as wc
    word = wc.Dispatch("Word.Application")
    word.Visible = False
    word.DisplayAlerts = False
    built = []
    try:
        for lang, manuscript, out_docx in jobs:
            num, front = build_toc.dump_pages(word, manuscript)
            build(lang, num, front, out_docx)
            built.append(out_docx)
    finally:
        try:
            word.Quit()
        except Exception:
            pass

    if not args.no_pdf:
        from docx2pdf import convert
        for out_docx in built:
            pdf = out_docx.with_suffix(".pdf")
            convert(str(out_docx), str(pdf))
            print("[pdf ]", pdf.name)


if __name__ == "__main__":
    main()
