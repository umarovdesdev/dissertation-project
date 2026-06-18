"""Build the GOST front-matter sections (EN + KZ) as .docx (+ .pdf):
NORMATIVE REFERENCES, DESIGNATIONS AND ABBREVIATIONS, DEFINITIONS.

Format follows the IITU house style observed in real council sample
dissertations (D:/dissertation_council/Образцы документов/авторы):
  * order — NORMATIVE REFERENCES, then DESIGNATIONS AND ABBREVIATIONS,
    then DEFINITIONS;
  * normative references — an intro line + each standard as a paragraph;
  * designations and abbreviations — a borderless two-column list
    (abbreviation | decoding), no header row;
  * definitions — flowing "Term — definition" paragraphs (bold term).

Normative refs and definitions are plain-paragraph Markdown, rendered with
md2gost.convert(). The abbreviations list needs a borderless, header-less
two-column table, so it is rendered here directly (md2gost only emits bordered
tables with a bold header row).

Sources : thesis/output/{normative_references,abbreviations,definitions}_{en,kz}.md
Output  : defense/docs/{NORMATIVE_REFERENCES,DESIGNATIONS_AND_ABBREVIATIONS,
          DEFINITIONS}_{EN,KZ}_GOST_<date>.docx (+ .pdf)

Usage:
    python build_frontmatter.py [--date YYYY-MM-DD] [--no-pdf]
"""
from __future__ import annotations

import argparse
import importlib.util
import re
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Mm, Pt

_HERE = Path(__file__).resolve().parent
ROOT = _HERE
while ROOT.parent != ROOT and not (ROOT / "defense").is_dir():
    ROOT = ROOT.parent

_spec = importlib.util.spec_from_file_location("md2gost", _HERE / "md2gost.py")
md2gost = importlib.util.module_from_spec(_spec)
_spec.loader.exec_module(md2gost)

_HDR = re.compile(r"^#{1,6}\s+(.*)$")
_ROW = re.compile(r"^\s*\|(.*)\|\s*$")


def _parse_row(line: str) -> list[str]:
    return [c.strip() for c in _ROW.match(line).group(1).split("|")]


def render_abbrev_into(doc, text: str) -> None:
    """Add a heading + borderless two-column list (no header row) to `doc`.

    Shared by `render_abbrev()` (standalone section) and the front-matter
    bundle, which composes this section into one combined document.
    """
    rows: list[list[str]] = []
    for raw in text.splitlines():
        s = raw.strip()
        if not s:
            continue
        m = _HDR.match(s)
        if m:
            md2gost._heading(doc, m.group(1).strip(), 1)
            continue
        if _ROW.match(s):
            cells = _parse_row(s)
            if set("".join(cells)) <= set("-: "):  # separator row, if any
                continue
            rows.append(cells[:2] if len(cells) >= 2 else [cells[0], ""])

    if rows:
        table = doc.add_table(rows=len(rows), cols=2)  # default style = no borders
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.allow_autofit = False
        for i, (abbr, dec) in enumerate(rows):
            for j, val, w in ((0, abbr, 32.0), (1, dec, 138.0)):
                cell = table.cell(i, j)
                cell.width = Mm(w)
                p = cell.paragraphs[0]
                p.paragraph_format.first_line_indent = Mm(0)
                p.paragraph_format.line_spacing = 1.0
                p.paragraph_format.space_after = Pt(0)
                md2gost._add_runs(p, val)


def render_abbrev(md_path: Path, docx_path: Path) -> None:
    """Standalone DESIGNATIONS AND ABBREVIATIONS section (.docx)."""
    text = md2gost.strip_version_markers(md_path.read_text(encoding="utf-8"))
    doc = Document()
    md2gost._configure_styles(doc)
    md2gost._configure_page(doc)
    render_abbrev_into(doc, text)
    md2gost._add_page_numbers(doc)
    docx_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(docx_path))
    print("[docx]", docx_path.name)


def main() -> None:
    ap = argparse.ArgumentParser(description="Build GOST front-matter (EN+KZ)")
    ap.add_argument("--date", default="2026-06-17")
    ap.add_argument("--no-pdf", action="store_true")
    args = ap.parse_args()

    src = ROOT / "thesis/output"
    docs = ROOT / "defense/docs"
    # (source stem, output name, renderer)
    jobs = [
        ("normative_references_en", "NORMATIVE_REFERENCES_EN", "para"),
        ("normative_references_kz", "NORMATIVE_REFERENCES_KZ", "para"),
        ("abbreviations_en", "DESIGNATIONS_AND_ABBREVIATIONS_EN", "abbrev"),
        ("abbreviations_kz", "DESIGNATIONS_AND_ABBREVIATIONS_KZ", "abbrev"),
        ("definitions_en", "DEFINITIONS_EN", "para"),
        ("definitions_kz", "DEFINITIONS_KZ", "para"),
    ]
    built = []
    for stem, name, kind in jobs:
        md = src / f"{stem}.md"
        docx = docs / f"{name}_GOST_{args.date}.docx"
        lang = "kz" if stem.endswith("_kz") else "en"
        if kind == "abbrev":
            render_abbrev(md, docx)
        else:
            md2gost.convert(md, docx, lang=lang)
            print("[docx]", docx.name)
        built.append(docx)

    if not args.no_pdf:
        from docx2pdf import convert as to_pdf
        for docx in built:
            pdf = docx.with_suffix(".pdf")
            to_pdf(str(docx), str(pdf))
            print("[pdf ]", pdf.name)


if __name__ == "__main__":
    main()
