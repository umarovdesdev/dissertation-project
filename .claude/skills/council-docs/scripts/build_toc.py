"""Build the GOST СОДЕРЖАНИЕ / CONTENTS document (EN + KZ) as .docx (+ .pdf).

Unlike the chapter body, a GOST table of contents needs real page numbers with
dotted leaders. This script reads the true pagination from the already-assembled
manuscripts in defense/docs/ (via MS Word over COM), maps each outline entry to
its page, and emits a standalone contents document with right-aligned dotted
leaders. Entries whose section is not yet written in the manuscript receive an
em-dash placeholder, so the contents stays honest to the current draft.

Outline sources : thesis/output/contents_en.md, thesis/output/contents_kz.md
Manuscripts     : defense/docs/DISSERTATION_{EN,KZ}_GOST_<date>.docx
Output          : defense/docs/TABLE_OF_CONTENTS_{EN,KZ}_GOST_<date>.docx (+ .pdf)

Usage:
    python build_toc.py [--date YYYY-MM-DD] [--no-pdf]
"""
from __future__ import annotations

import argparse
import importlib.util
import re
import time
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.shared import Mm, Pt

_HERE = Path(__file__).resolve().parent
# repo root: walk up until a dir containing defense/
ROOT = _HERE
while ROOT.parent != ROOT and not (ROOT / "defense").is_dir():
    ROOT = ROOT.parent

_spec = importlib.util.spec_from_file_location("md2gost", _HERE / "md2gost.py")
md2gost = importlib.util.module_from_spec(_spec)
_spec.loader.exec_module(md2gost)

TAB_MM = 170.0  # text-area right edge: A4 210 - left 30 - right 10
DASH = "—"
INDENT_MM = {1: 0.0, 2: 0.0, 3: 7.0, 4: 14.0}

_LEADNUM = re.compile(r"^§?\s*(\d+(?:\.[0-9A-Za-z]+)*)")
_CONCL_EN = re.compile(r"Conclusions to Chapter (\d+)", re.I)
_CONCL_KZ = re.compile(r"(\d+)-тарау")


# --- page-number extraction from the assembled manuscript ---------------------
def dump_pages(word, docx_path: Path) -> tuple[dict[str, int], dict[str, int]]:
    """Return (numbered, frontmatter) heading->page maps from a manuscript."""
    def _safe(fn):
        for _ in range(10):
            try:
                return fn()
            except Exception:
                time.sleep(0.4)
        return fn()

    import pythoncom  # noqa: F401  (ensures COM is initialised in this thread)

    doc = _safe(lambda: word.Documents.Open(str(docx_path), ReadOnly=True))
    fronts = (
        "INTRODUCTION", "CONCLUSION", "LIST OF REFERENCES", "REFERENCES",
        "APPENDIC", "NORMATIVE", "DEFINITION", "DESIGNATION",
        "КІРІСПЕ", "ҚОРЫТЫНДЫ", "ПАЙДАЛАН", "ҚОСЫМША", "НОРМАТИВ",
        "АНЫҚТАМА", "БЕЛГІЛЕУ",
    )
    num: dict[str, int] = {}
    front: dict[str, int] = {}
    for p in doc.Paragraphs:
        rng = p.Range
        b = rng.Bold
        if not (b is True or b == -1):
            continue
        t = rng.Text.strip().replace(chr(7), "").replace(chr(13), "")
        if not t:
            continue
        mm = _LEADNUM.match(t)
        page = int(rng.Information(3))  # wdActiveEndPageNumber
        if mm:
            num.setdefault(mm.group(1), page)
        elif any(t.upper().startswith(k) for k in fronts):
            front.setdefault(t.upper(), page)
    _safe(lambda: doc.Close(False))
    return num, front


def page_for_num(key: str, num: dict[str, int]):
    if key in num:
        return num[key]
    children = [p for k, p in num.items() if k.startswith(key + ".")]
    return min(children) if children else None


def resolve_page(text: str, num, front):
    mm = _LEADNUM.match(text)
    if mm:
        return page_for_num(mm.group(1), num)
    cm = _CONCL_EN.search(text) or _CONCL_KZ.search(text)
    if cm:
        return num.get(f"{cm.group(1)}.C")
    return front.get(text.strip().upper())


# --- outline parsing + document emission --------------------------------------
def parse_md(md_path: Path):
    for raw in md_path.read_text(encoding="utf-8").splitlines():
        s = raw.strip()
        if not s:
            continue
        if s.startswith("# "):
            yield (0, "title", s[2:].strip())
        elif s.startswith("## "):
            yield (1, "main", s[3:].strip())
        elif s.startswith("### "):
            yield (2, "entry", s[4:].strip())
        elif s.startswith("- "):
            txt = s[2:].strip()
            letters = [c for c in txt if c.isalpha()]
            is_caps = bool(letters) and all(c.isupper() for c in letters) and not txt[0].isdigit()
            mm = _LEADNUM.match(txt)
            if is_caps:
                yield (1, "main", txt)
            elif mm:
                yield (3 if mm.group(1).count(".") <= 2 else 4, "entry", txt)
            else:
                yield (3, "entry", txt)


def add_entry(doc, text, page, *, bold, indent_mm):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    pf.left_indent = Mm(indent_mm)
    pf.first_line_indent = Mm(0)
    pf.space_after = Pt(2)
    pf.line_spacing = 1.0
    pf.tab_stops.add_tab_stop(Mm(TAB_MM), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)
    md2gost._set_cell_font(p.add_run(text), bold=bold)
    md2gost._set_cell_font(p.add_run("\t" + (str(page) if page is not None else DASH)), bold=bold)


def build(md_path: Path, num, front, out_docx: Path):
    doc = Document()
    md2gost._configure_styles(doc)
    md2gost._configure_page(doc)
    for level, kind, text in parse_md(md_path):
        if kind == "title":
            h = doc.add_paragraph()
            h.alignment = WD_ALIGN_PARAGRAPH.CENTER
            h.paragraph_format.space_after = Pt(12)
            r = h.add_run(text)
            md2gost._set_cell_font(r, bold=True)
            r.font.size = Pt(16)
            continue
        add_entry(doc, text, resolve_page(text, num, front),
                  bold=(kind == "main"), indent_mm=INDENT_MM[level])
    md2gost._add_page_numbers(doc)
    out_docx.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_docx))
    print("[docx]", out_docx)


def main() -> None:
    ap = argparse.ArgumentParser(description="Build GOST CONTENTS (EN+KZ) with page numbers")
    ap.add_argument("--date", default="2026-06-17", help="manuscript/output date stamp")
    ap.add_argument("--no-pdf", action="store_true")
    args = ap.parse_args()

    docs = ROOT / "defense/docs"
    out_dir = ROOT / "thesis/output"
    jobs = [
        ("en", out_dir / "contents_en.md", docs / f"DISSERTATION_EN_GOST_{args.date}.docx",
         docs / f"TABLE_OF_CONTENTS_EN_GOST_{args.date}.docx"),
        ("kz", out_dir / "contents_kz.md", docs / f"DISSERTATION_KZ_GOST_{args.date}.docx",
         docs / f"TABLE_OF_CONTENTS_KZ_GOST_{args.date}.docx"),
    ]

    import win32com.client as wc
    word = wc.Dispatch("Word.Application")
    word.Visible = False
    word.DisplayAlerts = False
    built = []
    try:
        for _lang, md, manuscript, out_docx in jobs:
            num, front = dump_pages(word, manuscript)
            build(md, num, front, out_docx)
            built.append(out_docx)
    finally:
        try:
            word.Quit()
        except Exception:
            pass

    if not args.no_pdf:
        from docx2pdf import convert
        for out_docx in built:
            pdf = out_docx.with_suffix(".pdf")
            convert(str(out_docx), str(pdf))
            print("[pdf ]", pdf)


if __name__ == "__main__":
    main()
