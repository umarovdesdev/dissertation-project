"""Convert a council Markdown document into a GOST-formatted .docx (and optionally .pdf).

GOST parameters (per council/en/02-formatting/gost-formatting.md):
    A4, single line spacing, Times New Roman 14 pt,
    margins: left 30 mm, right 10 mm, top 20 mm, bottom 20 mm,
    page numbers centered at the bottom (not printed on the first page).

Markdown supported: # / ## / ### / #### headings, **bold**, *italic*, `code`,
numbered lists (1.), bullet lists (- / *), --- rule, blank-line paragraphs.

Usage:
    python md2gost.py INPUT.md [-o OUTPUT.docx] [--pdf]
"""

from __future__ import annotations

import argparse
import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt

FONT_NAME = "Times New Roman"
FONT_SIZE = 14  # pt
FIRST_LINE_INDENT_CM = 1.25

_INLINE = re.compile(r"(\*\*.+?\*\*|\*[^*].*?\*|`.+?`|\$[^$\n]+\$)")

# --- Version-marker scrubbing -------------------------------------------------
# Council deliverables are rendered OUTSIDE thesis/ (into defense/docs/). Per the
# project versioning policy, version markers — including the "V5" proper noun for
# the preprocessing pipeline — must never leak outside thesis/. The source .md in
# thesis/output/ is allowed to keep them; this converter strips them on the way
# out so the .docx/.pdf never carry a version. See REFACTORING.md §"V5 leak".

# Each pattern eats one adjacent space (leading where present) so removal leaves
# no double space and no space before punctuation — without touching whitespace
# elsewhere on the line (e.g. signature underscores spaced for layout).
# Parenthetical version tag, e.g. " (V5)", "(v5.1)", "(version 5.0)", "(нұсқа 5)".
_VER_PAREN = re.compile(
    r"[ \t]*\((?:[Vv][345](?:\.\d+)*|(?:version|версия|версии|нұсқа)[ \t]*[345](?:\.\d+)*)\)",
    re.IGNORECASE,
)
# Bare token, e.g. " V5", " v5.2", " V4.1", " V3" (leading space consumed).
_VER_TOKEN = re.compile(r"[ \t]*\b[Vv][345](?:\.\d+)*\b")
# Word form, e.g. " version 5.0", " версия 5", " нұсқа 5.1" (leading space consumed).
_VER_WORD = re.compile(
    r"[ \t]*\b(?:version|версия|версии|нұсқа)[ \t]*[345](?:\.\d+)*\b", re.IGNORECASE
)


def strip_version_markers(text: str) -> str:
    """Remove version markers (V3/V4/V5, decimals, word forms) from `text`.

    Council deliverables render outside thesis/, where no version marker —
    including the "V5" pipeline proper noun — may appear. Each pattern consumes
    the space preceding the marker so the surrounding text stays clean without
    rewriting unrelated whitespace.

    Args:
        text: Source Markdown that may legitimately contain version markers
            (it lives under thesis/).

    Returns:
        The text with version markers removed.
    """
    text = _VER_PAREN.sub("", text)
    text = _VER_TOKEN.sub("", text)
    text = _VER_WORD.sub("", text)
    return text


def _set_cell_font(run, *, bold=False, italic=False) -> None:
    run.font.name = FONT_NAME
    run.font.size = Pt(FONT_SIZE)
    # Ensure the font also applies to complex-script / Cyrillic ranges.
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rfonts.set(qn(attr), FONT_NAME)
    run.bold = bold
    run.italic = italic


# --- LaTeX math rendering -----------------------------------------------------
# The dissertation source carries math as LaTeX, inline (`$…$`) and display
# (`$$…$$`, optionally with `\tag{N}`). Word here has no equation engine, so this
# converter renders math as clean Unicode text with real super/subscript runs:
# `\beta\,A/L` → "βA/L", `$T/80$` → "T/80", `\frac{T}{80}` → "T/80",
# `p_t` → p with a subscript t, `A^k` → A with a superscript k. This removes the
# stray `$`, backslash commands and braces that otherwise leak into the .docx/.pdf.

_TEX_CMD = re.compile(r"\\([a-zA-Z]+)")

# Control symbols after a backslash (non-letter): spacing, escapes, line break.
_TEX_CTRL = {
    ",": "", "!": "", ";": " ", ":": " ", " ": " ", "\\": " ",
    "_": "_", "{": "{", "}": "}", "%": "%", "#": "#", "&": "&", "$": "$",
}

# Symbol commands → Unicode (Greek lower/upper, operators, relations, brackets).
_TEX_SYM = {
    # Greek (lowercase)
    "alpha": "α", "beta": "β", "gamma": "γ", "delta": "δ", "epsilon": "ε",
    "varepsilon": "ε", "zeta": "ζ", "eta": "η", "theta": "θ", "vartheta": "ϑ",
    "iota": "ι", "kappa": "κ", "lambda": "λ", "mu": "μ", "nu": "ν", "xi": "ξ",
    "pi": "π", "rho": "ρ", "sigma": "σ", "tau": "τ", "upsilon": "υ", "phi": "φ",
    "varphi": "φ", "chi": "χ", "psi": "ψ", "omega": "ω",
    # Greek (uppercase)
    "Gamma": "Γ", "Delta": "Δ", "Theta": "Θ", "Lambda": "Λ", "Xi": "Ξ",
    "Pi": "Π", "Sigma": "Σ", "Phi": "Φ", "Psi": "Ψ", "Omega": "Ω",
    # Operators / relations
    "cdot": "·", "times": "×", "div": "÷", "pm": "±", "mp": "∓", "ast": "∗",
    "ge": "≥", "geq": "≥", "le": "≤", "leq": "≤", "ne": "≠", "neq": "≠",
    "approx": "≈", "equiv": "≡", "sim": "∼", "propto": "∝",
    "cap": "∩", "cup": "∪", "subset": "⊂", "subseteq": "⊆", "supset": "⊃",
    "in": "∈", "notin": "∉", "forall": "∀", "exists": "∃",
    "sum": "Σ", "prod": "∏", "int": "∫", "partial": "∂", "nabla": "∇",
    "infty": "∞", "to": "→", "rightarrow": "→", "leftarrow": "←",
    "Rightarrow": "⇒", "leftrightarrow": "↔", "top": "⊤", "perp": "⊥",
    "angle": "∠", "lceil": "⌈", "rceil": "⌉", "lfloor": "⌊", "rfloor": "⌋",
    "cdots": "⋯", "ldots": "…", "dots": "…", "quad": "  ", "qquad": "    ",
}

# Operator/function names rendered upright (as the word itself).
_TEX_FUNC = {
    "min", "max", "log", "ln", "exp", "sin", "cos", "tan", "det", "lim",
    "deg", "gcd", "arg", "dim", "ker", "sup", "inf",
}

# One-argument commands whose content is rendered upright (drop the wrapper).
_TEX_WRAP = {
    "text", "mathrm", "mathbf", "mathit", "mathsf", "mathcal", "mathbb",
    "operatorname", "texttt", "boldsymbol", "mathtt", "textbf", "textit",
}

# Commands that render to nothing (delimiter sizing, style directives): the
# surrounding bracket/character they qualify is kept verbatim.
_TEX_DROP = {
    "left", "right", "big", "Big", "bigg", "Bigg", "bigl", "bigr",
    "Bigl", "Bigr", "biggl", "biggr", "Biggl", "Biggr", "displaystyle",
    "textstyle", "scriptstyle", "limits", "nolimits",
}


def _read_braces(s: str, i: int) -> tuple[str, int]:
    """Given s[i] == '{', return (inner_text, index_after_matching_'}')."""
    depth, j, n = 0, i, len(s)
    while j < n:
        if s[j] == "{":
            depth += 1
        elif s[j] == "}":
            depth -= 1
            if depth == 0:
                return s[i + 1:j], j + 1
        j += 1
    return s[i + 1:], n  # unbalanced: take the rest


def _read_script_arg(s: str, i: int) -> tuple[str, int]:
    """Read the argument of a `_`/`^` at index i: `{group}`, `\\command`, or one char."""
    n = len(s)
    if i >= n:
        return "", i
    if s[i] == "{":
        return _read_braces(s, i)
    if s[i] == "\\":
        m = _TEX_CMD.match(s, i)
        if m:
            return m.group(0), m.end()
        return s[i:i + 2], min(i + 2, n)
    return s[i], i + 1


def _tex_runs(s: str) -> list[tuple[str, str]]:
    """Parse a LaTeX math fragment into (text, script) runs.

    `script` is "" (baseline), "sub", or "sup". Commands, Greek letters and
    operators are mapped to Unicode; `\\frac{a}{b}` becomes "a/b"; sub/superscripts
    become tagged runs so the caller can apply real Word run formatting.
    """
    runs: list[tuple[str, str]] = []
    buf: list[str] = []
    n = len(s)

    def flush() -> None:
        if buf:
            runs.append(("".join(buf), ""))
            buf.clear()

    i = 0
    while i < n:
        c = s[i]
        if c == "\\":
            if i + 1 < n and not s[i + 1].isalpha():
                buf.append(_TEX_CTRL.get(s[i + 1], s[i + 1]))
                i += 2
                continue
            m = _TEX_CMD.match(s, i)
            name = m.group(1)
            i = m.end()
            if name in _TEX_DROP:
                pass  # sizing/style directive: render nothing, keep what follows
            elif name in _TEX_WRAP:
                inner, i = (_read_braces(s, i) if i < n and s[i] == "{" else ("", i))
                flush()
                runs.extend(_tex_runs(inner))
            elif name in ("frac", "dfrac", "tfrac"):
                a, b = "", ""
                if i < n and s[i] == "{":
                    a, i = _read_braces(s, i)
                while i < n and s[i] == " ":
                    i += 1
                if i < n and s[i] == "{":
                    b, i = _read_braces(s, i)
                flush()
                runs.extend(_frac_runs(a, b))
            elif name in _TEX_SYM:
                buf.append(_TEX_SYM[name])
            elif name in _TEX_FUNC:
                flush()
                runs.append((name, ""))
            else:
                buf.append(name)  # unknown: best-effort, drop the backslash
        elif c == "{":
            inner, i = _read_braces(s, i)
            flush()
            runs.extend(_tex_runs(inner))
        elif c == "}":
            i += 1  # stray closing brace
        elif c in "_^":
            script = "sub" if c == "_" else "sup"
            i += 1
            while i < n and s[i] == " ":
                i += 1
            arg, i = _read_script_arg(s, i)
            flush()
            runs.append((_tex_plain(arg), script))
        else:
            buf.append(c)
            i += 1
    flush()
    return runs


def _tex_plain(s: str) -> str:
    """Flatten a LaTeX fragment to plain text (used for script arguments/tags)."""
    return "".join(t for t, _ in _tex_runs(s))


def _frac_runs(a: str, b: str) -> list[tuple[str, str]]:
    """Render `\\frac{a}{b}` as a/b runs, parenthesising compound numerator/denominator."""
    return _maybe_paren(a) + [("/", "")] + _maybe_paren(b)


def _maybe_paren(group: str) -> list[tuple[str, str]]:
    runs = _tex_runs(group)
    plain = "".join(t for t, _ in runs).strip()
    if len(plain) > 1 and re.search(r"[+\-−·×/ ]", plain):
        return [("(", "")] + runs + [(")", "")]
    return runs


def _add_math_runs(paragraph, latex: str, *, bold=False, italic=False) -> None:
    """Render a LaTeX math fragment into `paragraph` as formatted runs."""
    for text, script in _tex_runs(latex):
        if not text:
            continue
        r = paragraph.add_run(text)
        _set_cell_font(r, bold=bold, italic=italic)
        if script == "sub":
            r.font.subscript = True
        elif script == "sup":
            r.font.superscript = True


def _split_tag(latex: str) -> tuple[str, str | None]:
    """Split a display equation's body from its `\\tag{…}` number, if any."""
    m = re.search(r"\\tag\{([^}]*)\}", latex)
    if not m:
        return latex, None
    return (latex[:m.start()] + latex[m.end():]).strip(), m.group(1)


def _add_equation(doc: Document, latex: str) -> None:
    """Render a display equation centred, with any `\\tag{}` number flush right (GOST)."""
    inner, tag = _split_tag(latex)
    usable = 170.0  # A4 text width: 210 − 30 (left) − 10 (right) mm
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.first_line_indent = Mm(0)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.tab_stops.add_tab_stop(Mm(usable / 2), WD_TAB_ALIGNMENT.CENTER)
    if tag:
        p.paragraph_format.tab_stops.add_tab_stop(Mm(usable), WD_TAB_ALIGNMENT.RIGHT)
    _set_cell_font(p.add_run("\t"))
    _add_math_runs(p, inner.strip())
    if tag:
        _set_cell_font(p.add_run("\t(" + _tex_plain(tag).strip() + ")"))


def _add_runs(paragraph, text: str, *, bold=False, italic=False) -> None:
    """Add inline-formatted runs (**bold**, *italic*, `code`, `$math$`) to a paragraph."""
    for token in _INLINE.split(text):
        if not token:
            continue
        if token.startswith("**") and token.endswith("**"):
            _set_cell_font(paragraph.add_run(token[2:-2]), bold=True, italic=italic)
        elif token.startswith("`") and token.endswith("`"):
            r = paragraph.add_run(token[1:-1])
            _set_cell_font(r, bold=bold, italic=italic)
            r.font.name = "Consolas"
            r._element.get_or_add_rPr().find(qn("w:rFonts")).set(qn("w:ascii"), "Consolas")
            r._element.get_or_add_rPr().find(qn("w:rFonts")).set(qn("w:hAnsi"), "Consolas")
        elif token.startswith("$") and token.endswith("$") and len(token) >= 2:
            _add_math_runs(paragraph, token[1:-1], bold=bold, italic=italic)
        elif token.startswith("*") and token.endswith("*"):
            _set_cell_font(paragraph.add_run(token[1:-1]), bold=bold, italic=True)
        else:
            _set_cell_font(paragraph.add_run(token), bold=bold, italic=italic)


def _configure_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = FONT_NAME
    normal.font.size = Pt(FONT_SIZE)
    normal.element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
    pf = normal.paragraph_format
    pf.line_spacing = 1.0
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)


def _configure_page(doc: Document) -> None:
    for section in doc.sections:
        section.page_width = Mm(210)
        section.page_height = Mm(297)
        section.left_margin = Mm(30)
        section.right_margin = Mm(10)
        section.top_margin = Mm(20)
        section.bottom_margin = Mm(20)
        section.different_first_page_header_footer = True  # no number on page 1


def _add_page_numbers(doc: Document) -> None:
    """Centered PAGE field in the footer; first-page footer left blank."""
    section = doc.sections[0]
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run = p.add_run()
    _set_cell_font(run)
    run._element.append(fld_begin)
    run._element.append(instr)
    run._element.append(fld_end)
    # Leave the first-page footer empty.
    section.first_page_footer.is_linked_to_previous = False


def _add_hrule(doc: Document) -> None:
    p = doc.add_paragraph()
    p_pr = p._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "auto")
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def _heading(doc: Document, text: str, level: int):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12 if level <= 2 else 6)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_with_next = True
    if level == 1:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _add_runs(p, text, bold=True)
        for r in p.runs:
            r.font.size = Pt(16)
    else:
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        _add_runs(p, text, bold=True, italic=(level >= 4))
    return p


def _body(doc: Document, text: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Mm(FIRST_LINE_INDENT_CM * 10)
    _add_runs(p, text)
    return p


def _list_item(doc: Document, marker: str, text: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.left_indent = Mm(12.5)
    p.paragraph_format.first_line_indent = Mm(-7.0)  # hanging
    _add_runs(p, f"{marker}\t{text}")
    return p


def _add_table(doc: Document, rows: list[list[str]]) -> None:
    """Render a Markdown pipe-table as a bordered Word table (TNR, header bold).

    `rows` is the list of cell-text rows (the `|---|` separator already removed).
    Backward-compatible: only invoked when the source contains pipe tables, which
    council deliverables do not, so their rendering is unaffected.
    """
    if not rows:
        return
    ncols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=ncols)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, row in enumerate(rows):
        for j in range(ncols):
            cell = table.cell(i, j)
            cell.paragraphs[0].text = ""  # clear default empty run
            p = cell.paragraphs[0]
            p.paragraph_format.first_line_indent = Mm(0)
            p.paragraph_format.line_spacing = 1.0
            text = row[j] if j < len(row) else ""
            _add_runs(p, text, bold=(i == 0))
    # spacing paragraph after the table
    doc.add_paragraph()


def _add_code_block(doc: Document, code_lines: list[str]) -> None:
    """Render a fenced code block as left-aligned monospace lines (no indent)."""
    for ln in code_lines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.first_line_indent = Mm(0)
        p.paragraph_format.line_spacing = 1.0
        r = p.add_run(ln if ln else " ")
        r.font.name = "Consolas"
        r.font.size = Pt(FONT_SIZE - 2)
        rpr = r._element.get_or_add_rPr()
        rfonts = rpr.find(qn("w:rFonts"))
        if rfonts is None:
            rfonts = OxmlElement("w:rFonts")
            rpr.append(rfonts)
        for attr in ("w:ascii", "w:hAnsi", "w:cs"):
            rfonts.set(qn(attr), "Consolas")


_NUM = re.compile(r"^(\d+)\.\s+(.*)$")
_BUL = re.compile(r"^[-*]\s+(.*)$")
_HDR = re.compile(r"^(#{1,6})\s+(.*)$")
_TBL_ROW = re.compile(r"^\s*\|.*\|\s*$")
_TBL_SEP = re.compile(r"^\s*\|?[\s:|-]*-[\s:|-]*\|?\s*$")
_FENCE = re.compile(r"^\s*```")

# --- Figure placeholders ------------------------------------------------------
# Drafts carry figures as inline text markers `[FIG-3.1: caption — path/img.png]`
# (and `[TAB-…]`). This converter resolves each to an embedded image with a GOST
# caption below it ("Figure N – Title" / KZ "Сурет N – Атауы"), replacing the
# inline marker with a cross-reference. `pre`("…in ") / `post`(" …суретінде")
# are absorbed so the reference reads naturally in either language.
_FIG = re.compile(
    r"(?P<pre>\b[Ii]n\s+)?"
    r"\[(?:FIG|FIGURE)-(?P<num>[0-9]+(?:\.[0-9]+)?):\s*(?P<body>[^\]]*)\]"
    r"(?P<post>\s+сурет\w*)?"
)
_DASH = re.compile(r"\s+[—–-]\s+")  # caption — target separator (em/en/hyphen)


def _png_size(p: Path):
    """Return (w, h) in pixels for a PNG, else None."""
    try:
        with open(p, "rb") as f:
            head = f.read(26)
        if head[:8] == b"\x89PNG\r\n\x1a\n" and head[12:16] == b"IHDR":
            return int.from_bytes(head[16:20], "big"), int.from_bytes(head[20:24], "big")
    except OSError:
        pass
    return None


def _fit_width_mm(p: Path, maxw: float = 165.0, maxh: float = 215.0) -> float:
    """Width (mm) that fits the image inside the text box, preserving aspect."""
    sz = _png_size(p)
    if not sz:
        return maxw * 0.9
    w, h = sz
    return min(maxw, maxh * w / h)


def _parse_fig_body(body: str, base: Path):
    """Split '[caption — target]' into (caption, resolved_image_path_or_None)."""
    segs = _DASH.split(body)
    if len(segs) >= 2:
        caption = " – ".join(s.strip() for s in segs[:-1]).strip()
        target = segs[-1].strip()
    else:
        caption, target = body.strip(), ""
    target = target.strip(" `").replace("/…/", "/").replace("…/", "").replace("/…", "")
    img = None
    if re.search(r"\.(png|jpe?g)$", target, re.I):
        cand = (base / target).resolve()
        if cand.is_file():
            img = cand
    return caption, img


def _insert_figure(doc: Document, label: str, num: str, caption: str, img: Path | None):
    if img is not None:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = Mm(0)
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.keep_with_next = True
        p.add_run().add_picture(str(img), width=Mm(_fit_width_mm(img)))
    c = doc.add_paragraph()
    c.alignment = WD_ALIGN_PARAGRAPH.CENTER
    c.paragraph_format.first_line_indent = Mm(0)
    c.paragraph_format.space_after = Pt(6)
    text = f"{label} {num} – {caption}"
    if img is None:
        text += " [ресурс дайындалуда]" if label == "Сурет" else " [asset to be created]"
    _add_runs(c, text)


def _parse_table_row(line: str) -> list[str]:
    s = line.strip()
    if s.startswith("|"):
        s = s[1:]
    if s.endswith("|"):
        s = s[:-1]
    return [c.strip() for c in s.split("|")]


def render_into(
    doc,
    text: str,
    *,
    lang: str = "en",
    base_dir: Path | None = None,
) -> None:
    """Render Markdown `text` into an existing (already-configured) document.

    Contains the full Markdown parsing loop but performs no page setup, footer,
    or save — so callers can compose several Markdown bodies into one document
    (e.g. the front-matter + manuscript bundle). `convert()` wraps this for the
    single-file case. Version-marker scrubbing is the caller's responsibility
    here (convert() still does it).
    """
    fig_label = "Сурет" if lang == "kz" else "Figure"
    if base_dir is None:
        base_dir = Path(".")
    lines = text.splitlines()

    figs: dict[str, dict] = {}  # num -> {caption, img, placed}

    def _fig_inline(m: re.Match) -> str:
        """Register the figure and return its inline cross-reference text."""
        num = m.group("num")
        if num not in figs:
            caption, img = _parse_fig_body(m.group("body"), base_dir)
            figs[num] = {"caption": caption, "img": img, "placed": False}
        if lang == "kz":
            post = (m.group("post") or "").strip()
            return f"{num}-{post}" if post else f"({num}-сурет)"
        pre = m.group("pre") or ""
        return f"{pre}Figure {num}" if pre else f"(Figure {num})"

    def _emit_figs(order: list[str]) -> None:
        for num in order:
            f = figs[num]
            if f["placed"]:
                continue
            _insert_figure(doc, fig_label, num, f["caption"], f["img"])
            f["placed"] = True

    buf: list[str] = []

    def flush_paragraph() -> None:
        if not buf:
            return
        raw = " ".join(buf).strip()
        buf.clear()
        order = [m.group("num") for m in _FIG.finditer(raw)]
        cleaned = _FIG.sub(_fig_inline, raw)
        _body(doc, cleaned)
        _emit_figs(order)

    tbl_buf: list[list[str]] = []

    def flush_table() -> None:
        if tbl_buf:
            _add_table(doc, tbl_buf)
            tbl_buf.clear()

    in_code = False
    code_buf: list[str] = []

    for raw in lines:
        line = raw.rstrip()
        stripped = line.strip()

        # fenced code block: collect raw lines verbatim until the closing fence
        if _FENCE.match(stripped):
            if in_code:
                _add_code_block(doc, code_buf)
                code_buf.clear()
                in_code = False
            else:
                flush_paragraph()
                flush_table()
                in_code = True
            continue
        if in_code:
            code_buf.append(raw.rstrip("\n"))
            continue

        # pipe table: accumulate consecutive table rows, skip the |---| separator
        if _TBL_ROW.match(line):
            flush_paragraph()
            if _TBL_SEP.match(line):
                continue
            tbl_buf.append(_parse_table_row(line))
            continue
        else:
            flush_table()

        if not stripped:
            flush_paragraph()
            continue
        if stripped == "---" or set(stripped) == {"-"} and len(stripped) >= 3:
            flush_paragraph()
            _add_hrule(doc)
            continue

        # display equation on its own line: $$ … $$  (optionally with \tag{N})
        if stripped.startswith("$$") and stripped.endswith("$$") and len(stripped) >= 4:
            flush_paragraph()
            flush_table()
            _add_equation(doc, stripped[2:-2])
            continue

        m = _HDR.match(stripped)
        if m:
            flush_paragraph()
            _heading(doc, m.group(2).strip(), len(m.group(1)))
            continue

        m = _NUM.match(stripped)
        if m:
            flush_paragraph()
            _list_item(doc, f"{m.group(1)}.", m.group(2).strip())
            continue

        m = _BUL.match(stripped)
        if m:
            flush_paragraph()
            _list_item(doc, "•", m.group(1).strip())
            continue

        buf.append(stripped)

    if in_code and code_buf:
        _add_code_block(doc, code_buf)
    flush_table()
    flush_paragraph()


def convert(
    md_path: Path,
    docx_path: Path,
    *,
    strip_versions: bool = True,
    lang: str | None = None,
    base_dir: Path | None = None,
) -> None:
    text = md_path.read_text(encoding="utf-8")
    if strip_versions:
        text = strip_version_markers(text)
    if lang is None:
        lang = "kz" if "_KZ_" in md_path.name else "en"
    if base_dir is None:  # repo root: walk up until a dir containing defense/
        base_dir = md_path.resolve().parent
        while base_dir.parent != base_dir and not (base_dir / "defense").is_dir():
            base_dir = base_dir.parent
    doc = Document()
    _configure_styles(doc)
    _configure_page(doc)
    render_into(doc, text, lang=lang, base_dir=base_dir)
    _add_page_numbers(doc)
    docx_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(docx_path))


def main() -> None:
    ap = argparse.ArgumentParser(description="Markdown -> GOST .docx (+ optional .pdf)")
    ap.add_argument("input", type=Path, help="input .md file")
    ap.add_argument("-o", "--output", type=Path, help="output .docx (default: alongside input)")
    ap.add_argument("--pdf", action="store_true", help="also render a .pdf via MS Word")
    args = ap.parse_args()

    md_path: Path = args.input
    docx_path: Path = args.output or md_path.with_suffix(".docx")
    convert(md_path, docx_path)
    print(f"[docx] {docx_path}")

    if args.pdf:
        from docx2pdf import convert as to_pdf

        pdf_path = docx_path.with_suffix(".pdf")
        to_pdf(str(docx_path), str(pdf_path))
        print(f"[pdf ] {pdf_path}")


if __name__ == "__main__":
    main()
