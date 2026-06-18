"""Build the GOST PhD dissertation TITLE PAGE (EN + KZ) as .docx (+ .pdf).

The title page is positional (centered organization block; UDC left / "as a
manuscript" right on one line; centered author, title, programme, degree
statement; a consultant block; centered place/year at the bottom), so it is
built directly with python-docx rather than through the Markdown converter. It
reuses md2gost's GOST page/style/font helpers so the page geometry (A4, TNR 14,
margins 30/10/20/20 mm) matches every other deliverable.

Required elements follow council/en/10-dissertation/structure.md §3.1.1:
organization, UDC, full name, title, programme code+name, sought degree,
scientific consultant, place/year.

Usage:
    python build_title.py [--date YYYY-MM-DD] [--no-pdf]
"""
from __future__ import annotations

import argparse
import importlib.util
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.shared import Mm, Pt

_HERE = Path(__file__).resolve().parent
ROOT = _HERE
while ROOT.parent != ROOT and not (ROOT / "defense").is_dir():
    ROOT = ROOT.parent

_spec = importlib.util.spec_from_file_location("md2gost", _HERE / "md2gost.py")
md2gost = importlib.util.module_from_spec(_spec)
_spec.loader.exec_module(md2gost)

TAB_MM = 170.0  # text-area right edge: A4 210 - left 30 - right 10

# --- field values (canonical; see PROJECT_MEMORY/people-and-identifiers.md) ----
FIELDS = {
    "en": {
        "org": [
            "International Information Technology University",
        ],
        "udc": "UDC: 004.93:617.735",
        "manuscript": "On manuscript right",
        "author": "YESMUKHAMEDOV NURMAGANBET SEITKALIULY",
        "title": "Automated Diabetic Retinopathy Diagnosis via Fundus Image "
                 "Enhancement and CNN Classification",
        "programme": "8D06104 – Computer systems and software engineering",
        "degree": [
            "Thesis for the degree of doctor of",
            "Philosophy (PhD)",
        ],
        "consultant": [
            ("Scientific consultant", False),
            ("Candidate of Phys.-Math. Sciences,", False),
            ("Associate Professor, International Information Technology University", False),
            ("Sapakova S.Z.", False),
            ("", False),
            ("Foreign consultant", False),
            ("Professor, Universiti Putra Malaysia", False),
            ("Al-Haddad S.A.R.", False),
        ],
        "place": ["Republic of Kazakhstan", "Almaty, 2026"],
    },
    "kz": {
        "org": [
            "Халықаралық ақпараттық технологиялар университеті",
        ],
        "udc": "ӘОЖ: 004.93:617.735",
        "manuscript": "Қолжазба құқығында",
        "author": "ЕСМУХАМЕДОВ НҰРМАҒАНБЕТ СЕЙТҚАЛИҰЛЫ",
        "title": "Fundus кескіндерін жақсарту және CNN жіктеуі арқылы "
                 "диабеттік ретинопатияны автоматтандырылған диагностикалау",
        "programme": "8D06104 – Есептеу техникасы және бағдарламалық жасақтама",
        "degree": [
            "Философия докторы (PhD) дәрежесін",
            "алуға арналған диссертация",
        ],
        "consultant": [
            ("Ғылыми консультанты", False),
            ("физика-математика ғылымдарының кандидаты,", False),
            ("қауымдастырылған профессор, Халықаралық ақпараттық технологиялар университеті", False),
            ("Сапакова С.З.", False),
            ("", False),
            ("Шетелдік ғылыми консультанты", False),
            ("профессор, Universiti Putra Malaysia", False),
            ("Al-Haddad S.A.R.", False),
        ],
        "place": ["Қазақстан Республикасы", "Алматы, 2026"],
    },
}


def _centered(doc, text, *, bold=False, size=14, space_before=0, space_after=0):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format
    pf.first_line_indent = Mm(0)
    pf.line_spacing = 1.0
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    r = p.add_run(text)
    md2gost._set_cell_font(r, bold=bold)
    r.font.size = Pt(size)
    return p


def _gap(doc, n):
    for _ in range(n):
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.0
        p.paragraph_format.first_line_indent = Mm(0)


def populate(doc, lang: str) -> None:
    """Add the title-page content to an existing (already-configured) document.

    Used both by `build()` (standalone title page) and by the front-matter
    bundle, which composes the title page as the first page of one document.
    """
    f = FIELDS[lang]

    # Organization block (top, centered, bold)
    for line in f["org"]:
        _centered(doc, line, bold=True, size=14, space_after=2)

    # UDC (left) / "as a manuscript" (right) on one line
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Mm(0)
    p.paragraph_format.line_spacing = 1.0
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.tab_stops.add_tab_stop(Mm(TAB_MM), WD_TAB_ALIGNMENT.RIGHT)
    md2gost._set_cell_font(p.add_run(f["udc"]))
    md2gost._set_cell_font(p.add_run("\t" + f["manuscript"]))

    _gap(doc, 7)
    _centered(doc, f["author"], bold=True, size=14)

    _gap(doc, 4)
    _centered(doc, f["title"], bold=True, size=16, space_after=6)
    _centered(doc, f["programme"], bold=False, size=14, space_before=12)

    _gap(doc, 1)
    for line in f["degree"]:
        _centered(doc, line, bold=False, size=14)

    # Consultant block — left-aligned, shifted into the right half of the page
    _gap(doc, 3)
    for text, bold in f["consultant"]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        pf = p.paragraph_format
        pf.left_indent = Mm(85)
        pf.first_line_indent = Mm(0)
        pf.line_spacing = 1.0
        pf.space_after = Pt(0)
        if text:
            md2gost._set_cell_font(p.add_run(text), bold=bold)

    # Place / year — centered, near the bottom
    _gap(doc, 6)
    for line in f["place"]:
        _centered(doc, line, bold=False, size=14)


def build(lang: str, out_docx: Path) -> None:
    doc = Document()
    md2gost._configure_styles(doc)
    md2gost._configure_page(doc)
    populate(doc, lang)
    out_docx.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_docx))
    print("[docx]", out_docx.name)


def main() -> None:
    ap = argparse.ArgumentParser(description="Build GOST TITLE PAGE (EN+KZ)")
    ap.add_argument("--date", default="2026-06-17", help="output date stamp")
    ap.add_argument("--no-pdf", action="store_true")
    args = ap.parse_args()

    docs = ROOT / "defense/docs"
    built = []
    for lang in ("en", "kz"):
        out = docs / f"TITLE_PAGE_{lang.upper()}_GOST_{args.date}.docx"
        build(lang, out)
        built.append(out)

    if not args.no_pdf:
        from docx2pdf import convert
        for out_docx in built:
            pdf = out_docx.with_suffix(".pdf")
            convert(str(out_docx), str(pdf))
            print("[pdf ]", pdf.name)


if __name__ == "__main__":
    main()
